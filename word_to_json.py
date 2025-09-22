#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文件读取器 - 将Word文档内容转换为JSON格式
支持.docx和.doc格式的Word文件
"""

import json
import os
import sys
from datetime import datetime
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    print("请安装python-docx库: pip install python-docx")
    sys.exit(1)

try:
    import win32com.client
    WORD_COM_AVAILABLE = True
except ImportError:
    WORD_COM_AVAILABLE = False
    print("警告: 无法导入win32com.client，将无法处理.doc文件")


class WordToJsonConverter:
    """Word文档转JSON转换器"""
    
    def __init__(self):
        self.supported_extensions = ['.docx', '.doc']
    
    def read_docx_file(self, file_path):
        """读取.docx文件"""
        try:
            doc = Document(file_path)
            content = {
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "file_path": str(file_path),
                    "file_size": os.path.getsize(file_path),
                    "processed_time": datetime.now().isoformat()
                },
                "document_structure": {
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(doc.tables),
                    "total_sections": len(doc.sections)
                },
                "content": {
                    "paragraphs": [],
                    "tables": [],
                    "headers_footers": []
                }
            }
            
            # 读取段落内容
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():  # 只保存非空段落
                    para_info = {
                        "index": i,
                        "text": paragraph.text,
                        "style": paragraph.style.name if paragraph.style else "Normal",
                        "alignment": self._get_alignment_text(paragraph.alignment),
                        "runs": []
                    }
                    
                    # 读取段落中的运行（格式化文本）
                    for run in paragraph.runs:
                        if run.text.strip():
                            run_info = {
                                "text": run.text,
                                "bold": run.bold,
                                "italic": run.italic,
                                "underline": run.underline,
                                "font_name": run.font.name,
                                "font_size": str(run.font.size) if run.font.size else None
                            }
                            para_info["runs"].append(run_info)
                    
                    content["content"]["paragraphs"].append(para_info)
            
            # 读取表格内容
            for i, table in enumerate(doc.tables):
                table_info = {
                    "index": i,
                    "rows": len(table.rows),
                    "columns": len(table.columns),
                    "data": []
                }
                
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell_idx, cell in enumerate(row.cells):
                        cell_data = {
                            "row": row_idx,
                            "column": cell_idx,
                            "text": cell.text.strip()
                        }
                        row_data.append(cell_data)
                    table_info["data"].append(row_data)
                
                content["content"]["tables"].append(table_info)
            
            # 读取页眉页脚
            for section in doc.sections:
                if section.header:
                    for paragraph in section.header.paragraphs:
                        if paragraph.text.strip():
                            content["content"]["headers_footers"].append({
                                "type": "header",
                                "text": paragraph.text
                            })
                
                if section.footer:
                    for paragraph in section.footer.paragraphs:
                        if paragraph.text.strip():
                            content["content"]["headers_footers"].append({
                                "type": "footer",
                                "text": paragraph.text
                            })
            
            return content
            
        except Exception as e:
            return {"error": f"读取.docx文件时出错: {str(e)}"}
    
    def read_doc_file(self, file_path):
        """读取.doc文件（需要win32com）"""
        if not WORD_COM_AVAILABLE:
            return {"error": "无法处理.doc文件，请安装pywin32库或将文件转换为.docx格式"}
        
        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            
            doc = word_app.Documents.Open(str(file_path))
            
            content = {
                "file_info": {
                    "filename": os.path.basename(file_path),
                    "file_path": str(file_path),
                    "file_size": os.path.getsize(file_path),
                    "processed_time": datetime.now().isoformat()
                },
                "document_structure": {
                    "total_paragraphs": doc.Paragraphs.Count,
                    "total_tables": doc.Tables.Count
                },
                "content": {
                    "paragraphs": [],
                    "tables": []
                }
            }
            
            # 读取段落
            for i in range(1, doc.Paragraphs.Count + 1):
                paragraph = doc.Paragraphs(i)
                if paragraph.Range.Text.strip():
                    para_info = {
                        "index": i - 1,
                        "text": paragraph.Range.Text.strip(),
                        "style": paragraph.Style.NameLocal
                    }
                    content["content"]["paragraphs"].append(para_info)
            
            # 读取表格
            for i in range(1, doc.Tables.Count + 1):
                table = doc.Tables(i)
                table_info = {
                    "index": i - 1,
                    "rows": table.Rows.Count,
                    "columns": table.Columns.Count,
                    "data": []
                }
                
                for row_idx in range(1, table.Rows.Count + 1):
                    row_data = []
                    for col_idx in range(1, table.Columns.Count + 1):
                        try:
                            cell_text = table.Cell(row_idx, col_idx).Range.Text.strip()
                            cell_data = {
                                "row": row_idx - 1,
                                "column": col_idx - 1,
                                "text": cell_text
                            }
                            row_data.append(cell_data)
                        except:
                            # 处理合并单元格的情况
                            row_data.append({
                                "row": row_idx - 1,
                                "column": col_idx - 1,
                                "text": ""
                            })
                    table_info["data"].append(row_data)
                
                content["content"]["tables"].append(table_info)
            
            doc.Close()
            word_app.Quit()
            
            return content
            
        except Exception as e:
            try:
                word_app.Quit()
            except:
                pass
            return {"error": f"读取.doc文件时出错: {str(e)}"}
    
    def _get_alignment_text(self, alignment):
        """获取对齐方式的文本描述"""
        alignment_map = {
            WD_PARAGRAPH_ALIGNMENT.LEFT: "left",
            WD_PARAGRAPH_ALIGNMENT.CENTER: "center",
            WD_PARAGRAPH_ALIGNMENT.RIGHT: "right",
            WD_PARAGRAPH_ALIGNMENT.JUSTIFY: "justify"
        }
        return alignment_map.get(alignment, "unknown")
    
    def convert_to_json(self, file_path, output_path=None, pretty_print=True):
        """将Word文件转换为JSON"""
        file_path = Path(file_path)
        
        if not file_path.exists():
            return {"error": f"文件不存在: {file_path}"}
        
        if file_path.suffix.lower() not in self.supported_extensions:
            return {"error": f"不支持的文件格式: {file_path.suffix}"}
        
        # 根据文件扩展名选择读取方法
        if file_path.suffix.lower() == '.docx':
            content = self.read_docx_file(file_path)
        else:  # .doc
            content = self.read_doc_file(file_path)
        
        # 输出JSON
        if output_path:
            try:
                with open(output_path, 'w', encoding='utf-8') as f:
                    if pretty_print:
                        json.dump(content, f, ensure_ascii=False, indent=2)
                    else:
                        json.dump(content, f, ensure_ascii=False)
                print(f"JSON文件已保存到: {output_path}")
            except Exception as e:
                print(f"保存JSON文件时出错: {e}")
        else:
            # 直接输出到控制台 - 已禁用详细输出
            pass
        
        return content


def main():
    """主函数 - 命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(
        description="将Word文件转换为JSON格式",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例用法:
  python word_to_json.py document.docx                    # 输出到控制台
  python word_to_json.py document.docx output.json       # 保存到文件
  python word_to_json.py --help                          # 显示帮助信息
        """
    )
    
    parser.add_argument(
        'word_file',
        help='要转换的Word文件路径 (.docx 或 .doc)'
    )
    
    parser.add_argument(
        'output_file',
        nargs='?',
        help='输出JSON文件路径 (可选，不指定则输出到控制台)'
    )
    
    parser.add_argument(
        '--version',
        action='version',
        version='Word to JSON Converter 1.0.0'
    )
    
    args = parser.parse_args()
    
    print(f"正在处理文件: {args.word_file}")
    
    converter = WordToJsonConverter()
    result = converter.convert_to_json(args.word_file, args.output_file)
    
    if "error" in result:
        print(f"错误: {result['error']}")
        sys.exit(1)
    else:
        if args.output_file:
            print(f"转换完成! JSON文件已保存到: {args.output_file}")
        else:
            print("转换完成! JSON内容已生成")


if __name__ == "__main__":
    main()